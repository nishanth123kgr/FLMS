from random import randint

import pandas as pd
from num2words import num2words
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfFileWriter, PdfReader, PdfWriter


def encrypt_pdf(pdf, password="!@#$%^&*()"):
    output_encrypted_pdf = "salary_reports/pdf/encrypted/" + pdf.split("/")[-1][:-4] + "_encrypted.pdf"

    output_pdf_file = open(pdf, "rb")
    output_pdf_reader = PdfReader(output_pdf_file)
    output_pdf_writer = PdfWriter()

    for page_num in range(len(output_pdf_reader.pages)):
        output_pdf_writer.add_page(output_pdf_reader.pages[page_num])

    output_pdf_writer.encrypt(password)

    with open(output_encrypted_pdf, "wb") as output_encrypted_pdf_file:
        output_pdf_writer.write(output_encrypted_pdf_file)
    output_pdf_file.close()
    print("Encrypted PDF created successfully!")



def generate_password(name):
    password = f"{randint(100000, 999999)}"
    print(password, name)
    return password


def gen_sal_report(row):
    doc = Document("leaveFormTemplates/sal_template.docx")
    placeholders = {
        '#name': row[0].strip(),
        '#des': row[1],
        '#bp': row[2],
        '#ddd': row[3],
        '#hra': row[4],
        '#cca': row[5],
        '#ma': row[6],
        '#te': row[7],
        '#cps': row[8],
        '#ca': row[9],
        '#it': row[10],
        '#nhis': row[11],
        '#fsf': row[12],
        '#fa': row[13],
        '#pt': row[14],
        '#mcp': row[15],
        '#td': row[16],
        '#np': f"Rs.{row[17]}/- (Rupees {num2words(row[17]).replace(',', '').replace(' and', '').replace('-', ' ').title()} only)"
    }
    for table in doc.tables:
        fixed_elems = []
        for rows in table.rows:
            for cell in rows.cells:
                if cell.text not in fixed_elems:
                    # if cell.text in placeholders:
                    #     print(cell.text, placeholders[cell.text])
                    fixed_elems.append(cell.text)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            if text in placeholders:
                                # print(text, placeholders[text])
                                run.text = placeholders[text]
    doc.save(f"salary_reports/{row[0].replace(' ', '')}.docx")
    pdf_loc = f"salary_reports/pdf/{row[0].replace(' ', '')}.pdf"
    convert(f"salary_reports/{row[0].replace(' ', '')}.docx", pdf_loc)
    encrypt_pdf(f"salary_reports/pdf/{row[0].replace(' ', '')}.pdf", generate_password(row[0]))


if __name__ == '__main__':
    data = pd.read_excel('Salary for the month of March 2023 (Scale) - General Fund (Scale).xls')
    df = data.iloc[2:, 1:]
    df.columns = df.iloc[0]
    df = df[1:]
    df = df.dropna()
    df.reset_index(drop=True, inplace=True)
    for index, row in df.iterrows():
        row = row.astype(str)
        row = row.to_list()

        gen_sal_report(row)
