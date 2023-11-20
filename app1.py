from flask import Flask, request, jsonify, render_template, send_file
import mysql.connector
from flask_cors import CORS
from datetime import datetime, timedelta
import docx
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
import random
import string
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

app = Flask(__name__)
CORS(app)

db = mysql.connector.connect(host="localhost", port="3306", user="root", database="facultyleavedb")
cursor = db.cursor()

def get_staff_from_name(name):
    cursor.execute(f'''
    SELECT id FROM staff
    WHERE name = '{name}'
    ''')
    namee = cursor.fetchone()
    cursor.reset()
    return jsonify(namee)

def set_run_font(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = font_size

def set_paragraph_alignment(paragraph, alignment):
    paragraph.alignment = alignment

def replace_text_in_docx(input_docx_file, output_docx_file, replacements, data_array):
    doc = docx.Document(input_docx_file)
    if len(doc.tables) > 0:
        table = doc.tables[0]
        for row_data in data_array:
            new_row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(new_row):
                    cell = new_row[i]
                    cell.text = str(cell_data)
                    for paragraph in cell.paragraphs:
                        set_run_font(paragraph.runs[0], "Arial", Pt(10))
                        set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)
    for field, value in replacements.items():
        replace_placeholder(doc, field, value)
    doc.save(output_docx_file)

def replace_placeholder(doc, target, replacement):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if target in run.text and replacement:
                new_text = run.text.replace(target, replacement)
                run.text = new_text

def generate_random_string(length=5):
    characters = string.ascii_letters + string.digits
    random_string = ''.join(random.choice(characters) for _ in range(length))
    return random_string

def insert_into_naan_mudhalvan_day(data):
    insert_query = """
    INSERT INTO naan_mudhalvan_day (department, semester, academic_year, academic_semester, day)
    VALUES (%s, %s, %s, %s, %s)
    """
    cursor.execute(insert_query, data)
    cursor.reset()
    db.commit()

def insert_into_timetable(data):
    insert_query = """
    INSERT INTO timetable (day, period, subject_id, academic_year, academic_semester)
    VALUES (%s, %s, %s, %s, %s)
    """

    cursor.executemany(insert_query, data)
    cursor.reset()
    db.commit()

def retrive_class_for_staff_in_specific_day(staff, day):
    query = f"""
    SELECT tt.period, s.id, s.semester, s.department              
    FROM timetable t
    INNER JOIN timetable_timing tt ON t.period = tt.period
    INNER JOIN subjects s ON t.subject_id = s.id
    INNER JOIN staff_handling sh ON s.id = sh.subject_id
    INNER JOIN staff st ON sh.staff_id = st.id
    WHERE st.id = {staff} AND t.day = '{day}';
    """

    cursor.execute(query)
    classes = cursor.fetchall()
    return classes

def get_weekdays_in_date_range(start_date_str, end_date_str):
    weekdays = ["MON", "TUE", "WED", "THU", "FRI"]
    start_date = datetime.strptime(start_date_str, "%d-%m-%Y")
    end_date = datetime.strptime(end_date_str, "%d-%m-%Y")
    current_date = start_date
    result = []

    while current_date <= end_date:
        if current_date.weekday() < 5:
            day_of_week = weekdays[current_date.weekday()]
            result.append((day_of_week, current_date.strftime("%d-%m-%Y")))

        current_date += timedelta(days=1)

    return result

@app.route('/api/add_timetable', methods=['POST'])
def api_add_timetable():
    dep = request.form['dep']
    sem = request.form['sem']
    acayear = request.form['acayear']
    acasem = request.form['acasem']
    naan = request.form['naan']

    naan_mudhalvan_data = (dep, sem, acayear, acasem, naan)

    timetable_data = []
    for day in ['mon', 'tue', 'wed', 'thu', 'fri']:
        for period in range(1, 9):
            subject_id = request.form[f'{day}_{period}']
            if subject_id == '000000':
                continue
            else:
                timetable_data.append((day.upper(), period, subject_id, acayear, acasem))

    insert_into_naan_mudhalvan_day(naan_mudhalvan_data)
    insert_into_timetable(timetable_data)

    return "Hello"
@app.route('/add_timetable', methods=['GET'])
def add_timetable():
    return render_template('add_timetable.html')

@app.route('/api/update_timetable_timing', methods=['POST'])
def update_timetable_timing():
    for period in range(1, 9):
        start_time = request.form.get(f'start_{period}')
        end_time = request.form.get(f'end_{period}')

        update_query = f"UPDATE `timetable_timing` SET `start_time`='{start_time}',`end_time`='{end_time}' WHERE  `period`='{period}''"
        cursor.execute(update_query)
        cursor.reset()
        db.commit()
    return "True"

@app.route('/update_timetable_timing', methods=['GET'])
def update_timetable_timing_():
    return render_template('update_timetable_timing.html')

@app.route('/api/insert_subjects', methods=['POST'])
def insert_subject():
    code = request.form.get("code")
    name = request.form.get("name")
    type = request.form.get("type")
    sem = request.form.get("sem")
    credit = request.form.get("credit")
    contact = request.form.get("contact")
    regu = request.form.get("regu")
    department = request.form.get("department")
    staff = request.form.get("staff")

    update_query = f"""
     INSERT 
     INTO `subjects`(`id`, `name`, `type`, `semester`, `credit_hour`, `contact_period`, `regulation`, `department`) 
     VALUES ('{code}', '{name}', '{type}', '{sem}', '{credit}', '{contact}', '{regu}', '{department}');
    """

    cursor.execute(update_query)
    cursor.reset()

    update_query = f"INSERT INTO `staff_handling`(`subject_id`, `staff_id`) VALUES ('{code}','{staff}')"
    cursor.execute(update_query)
    cursor.reset()
    db.commit()

    return "Subject updated successfully!"

@app.route('/add_subjects', methods=['GET'])
def add_subject():
    return render_template('add_subject.html')

@app.route("/api/retrive_all_staff", methods=["GET"])
def retrive_all_staff():
    cursor.execute('SELECT `id`,`name` FROM `staff`')
    staffs = cursor.fetchall()
    return jsonify(staffs)

@app.route("/api/retrive_all_subjects", methods=["GET"])
def retrive_all_subjects():
    cursor.execute('SELECT `id`,`name` FROM `subjects`')
    subjects = cursor.fetchall()
    return jsonify(subjects)

@app.route("/api/get_staffs_classes/<staff>/<fdate>/<tdate>", methods=["GET"])
def get_staffs_classes(staff, fdate, tdate):
    days = get_weekdays_in_date_range(fdate, tdate)
    staffs_classes = []

    for day in days:
        day_name, date = day[0], day[1]
        classes = retrive_class_for_staff_in_specific_day(staff, day_name)
        for clas in classes:
            staffs_classes.append([day_name, date, clas[0], clas[1], clas[2], clas[3]])
    return jsonify(staffs_classes)

@app.route("/api/get_free_staffs/<day>/<period>/<semester>/<department>", methods=["GET"])
def get_free_staffs(day, period, semester, department):
    cursor.execute('''SELECT
    tt.day AS timetable_day,
    tt.period AS timetable_period,
    s.id AS subject_id,
    s.name AS subject_name,
    s.semester AS subject_semester,
    s.department AS subject_department,
    sh.staff_id,
    st.name AS staff_name
FROM
    timetable tt
JOIN
    subjects s ON tt.subject_id = s.id
JOIN
    staff_handling sh ON tt.subject_id = sh.subject_id
JOIN
    staff st ON sh.staff_id = st.id;''')

    data = get_free_staff(cursor.fetchall(), str(department), str(day), str(period), str(semester))

    cursor.reset()
    return jsonify(data)

def get_free_staff(timetable, department, day, period, semester):
    staff_with_class = set()
    for row in timetable:
        if str(row[0]) == day and str(row[1]) == period:
            staff_with_class.add((row[6], row[7]))

    staff_in_department = set()
    for row in timetable:
        if str(row[5]) == department and str(row[4]) == semester:
            staff_in_department.add((row[6], row[7]))

    free_staff = staff_in_department - staff_with_class

    return list(free_staff)

@app.route('/api/get_staffs_subjects/<staff>/<department>/<semester>', methods=['GET'])
def get_staffs_subjects(staff, department, semester):
    query = f"""SELECT
    s.id AS subject_code,
    s.name AS subject_name
FROM
    staff_handling sh
JOIN
    subjects s ON sh.subject_id = s.id
JOIN
    staff st ON sh.staff_id = st.id
WHERE
    st.id = {staff}
    AND s.department = '{department}'
    AND s.semester = {semester}"""

    cursor.execute(query)

    data = cursor.fetchall()

    cursor.reset()

    return jsonify(data)

@app.route("/api/get_staff_id_from_name/<name>", methods=["GET"])
def get_staff_id_from_name(name):
    cursor.execute(f'''
    SELECT id FROM staff
    WHERE name = '{name}'
    ''')

    namee = cursor.fetchone()

    cursor.reset()
    return jsonify(namee)

@app.route("/api/changeStaff/<department>/<name>", methods=["GET"])
def changeStaff(department, name):
    query = f'''SELECT staff_handling.subject_id, subjects.name
FROM staff_handling
JOIN subjects ON subjects.id = staff_handling.subject_id
JOIN staff ON staff.id = staff_handling.staff_id
WHERE staff.name = '{name}'
AND subjects.department = '{department}';'''
    cursor.execute(query)
    subjects = cursor.fetchall()
    cursor.reset()
    return jsonify(subjects)

@app.route("/api/generate/el", methods=["POST"])
def genRate_el():
    data = request.json
    object_data = data.get('objectData', {})
    array_data = data.get('leaves', [])
    name = generate_random_string()
    name = name + '.docx'
    temp_dir = 'temp/'
    os.makedirs(temp_dir, exist_ok=True)
    file_path = os.path.join(temp_dir, name)
    replace_text_in_docx('leaveFormTemplates/el.docx', file_path, object_data, array_data)
    convert(file_path, file_path.replace('.docx', '.pdf'))
    return jsonify(file_path.replace('.docx', '.pdf'))

@app.route("/api/generate/cl", methods=["POST"])
def genRate_cl():
    data = request.json
    object_data = data.get('objectData', {})
    array_data = data.get('leaves', [])
    name = generate_random_string()
    name = name + '.docx'
    temp_dir = 'temp/'
    os.makedirs(temp_dir, exist_ok=True)
    file_path = os.path.join(temp_dir, name)
    replace_text_in_docx('leaveFormTemplates/cl.docx', file_path, object_data, array_data)
    convert(file_path, file_path.replace('.docx', '.pdf'))
    os.remove(file_path)
    return jsonify(file_path.replace('.docx', '.pdf'))

@app.route('/leave_form', methods=['GET'])
def leave_form():
    return render_template('leave_form.html')

@app.route('/api/download/<folder>/<file>', methods=['GET'])
def download(folder, file):
    file = f"{folder}/{file}"
    return send_file(file, as_attachment=True)

if __name__ == '__main__':
    app.run()
