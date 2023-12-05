import ast
import datetime
import hashlib
import json
import os
import random
import smtplib
import string
import urllib.parse
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import mysql.connector
import pandas as pd
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    flash,
    jsonify,
    session,
    send_file,
)
from flask_cors import CORS

from from_csv.el_from_csv import get_el
from from_csv.ml_mtl_lop_from_csv import get_ml_mtl_lop
from from_csv.utils import get_total_days, convert_to_date_text
from from_csv.vl_from_csv import get_vl

from utils.report_generation import generate_report
from utils.leave_calc import get_el_earned

import docx
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

app = Flask(__name__)
CORS(app)


db = mysql.connector.connect(
    host="localhost", port="3306", user="root", database="facultyleavedb"
)
cursor = db.cursor()

drive_directory = "H:\\My Drive\\data\\"
drive_folder = "H:\My Drive\data"

app.config["SECRET_KEY"] = "f6f5372c8210fd221e2b4842e51af432"
app.config["DRIVE_DIRECTORIES"] = {
    "CSE": drive_directory + "CSE",
    "ECE": drive_directory + "ECE",
    "MECH": drive_directory + "MECH",
    "CIVIL": drive_directory + "CIVIL",
    "MBA": drive_directory + "MBA",
    "MCA": drive_directory + "MCA",
    "SH": drive_directory + "SH",
}


def get_data(id, type):
    if session["role"] != 4 and session["role"] != 0:
        return redirect("/login")
    cursor.execute(
        "select * from %s where id = %s"
        % (
            type,
            id,
        )
    )
    result = cursor.fetchall()
    for i in range(len(result)):
        result[i] = list(result[i])
        del result[i][-1]
    cursor.reset()
    return result


@app.route("/")
def display_staff():
    # if session['role'] != 4:
    #     return redirect("/login")
    return redirect("/login")
    # cursor.execute("select * from staff")
    # result = cursor.fetchall()
    # cursor.reset()
    # return render_template("index.html", table_content=enumerate(result))


@app.route("/staff/<id>", methods=["POST"])
def get_staff_details(id):
    if session["role"] != 4:
        return redirect("/login")
    cursor.execute("select name, department, doj from staff where id = %s" % (id,))
    result = cursor.fetchall()
    cursor.reset()
    return jsonify(result)


@app.route("/dashboard")
def display_dashboard():
    session_key = session.get("session_key")
    role = str(session.get("role"))
    print(role, session_key)
    if session_key and role:
        if role == "4":
            cursor.execute(
                "SELECT name FROM `staff` WHERE id=(SELECT id from user_login where SESSION_key=%s)",
                (session_key,),
            )
            name = cursor.fetchone()[0]
            cursor.reset()
            return render_template("dashboard.html", name=name)
        elif role == "0":
            cursor.execute(
                "SELECT name, id, department FROM `staff` WHERE id=(SELECT id from user_login where SESSION_key=%s)",
                (session_key,),
            )
            data = cursor.fetchone()
            cursor.reset()
            return render_template(
                "staff_dashboard.html", name=data[0], id=data[1], dept=data[2]
            )
        else:
            return redirect("/login")
    else:
        return redirect("/login")


@app.route("/staff/upload", methods=["GET", "POST"])
def upload_file():
    if session["role"] != 4:
        return redirect("/login")
    if request.method == "GET":
        return render_template("upload.html")
    elif request.method == "POST":
        uploaded_file = request.files["file"]
        if uploaded_file:
            # Save the uploaded file to a desired location
            info = json.loads(request.form["info"])
            el = get_el(uploaded_file, info["sheetName"])
            ml, mtl, lop = get_ml_mtl_lop(uploaded_file, info["sheetName"])
            # Update the EL table
            for i in range(len(el)):
                row = el[i]
                query = ''
                try:
                    total = get_total_days(row[0], row[1])
                    query = 'INSERT INTO el (id, dept, `from`, `to`, from_prefix, to_prefix, from_suffix, to_suffix, date_of_rejoin, total) VALUES (%s, "%s", "%s", "%s", %s, %s, %s, %s, %s, %s)' % (
                        info["id"],
                        info["department"],
                        row[0],
                        row[1],
                        row[2][0] if row[2][0] == "NULL" else '"' + row[2][0] + '"',
                        row[2][1] if row[2][1] == "NULL" else '"' + row[2][1] + '"',
                        row[2][2] if row[2][2] == "NULL" else '"' + row[2][2] + '"',
                        row[2][3] if row[2][3] == "NULL" else '"' + row[2][3] + '"',
                        row[3] if row[3] == "NULL" else '"' + row[3] + '"',
                        total,
                    )
                    # print(query)
                    cursor.execute(
                        query
                    )
                    db.commit()
                    cursor.reset()
                except Exception as e:
                    if "Duplicate entry" not in str(e):
                        # print(query)
                        # print(row)
                        # print(e)
                        return {
                            "error": "Error parsing data on row "
                            + str(i + 12)
                            + " of EL sheet.\nError Description: "
                            + str(e)
                        }
            # Update the ML, MTL, LOP table
            for name, table in [("ml", ml), ("mtl", mtl), ("lop", lop)]:
                if table:
                    for i in range(len(table)):
                        row = table[i]
                        print(row)
                        try:
                            total = get_total_days(row[0], row[1])
                            query = 'INSERT INTO %s (id, dept, `from`, `to`, medical_fittness_on, from_prefix, to_prefix, from_suffix, to_suffix, doj,total) VALUES (%s, "%s", "%s", "%s",%s, %s, %s, %s, %s, "%s", %s)'% (
                                    name,
                                    info["id"],
                                    info["department"],
                                    row[0],
                                    row[1],
                                    "NULL" if (row[3] == "NULL" or row[3] == '-') else row[3],
                                    row[4][0]
                                    if row[4][0] == "NULL"
                                    else '"' + row[4][0] + '"',
                                    row[4][1]
                                    if row[4][1] == "NULL"
                                    else '"' + row[4][1] + '"',
                                    row[4][2]
                                    if row[4][2] == "NULL"
                                    else '"' + row[4][2] + '"',
                                    row[4][3]
                                    if row[4][3] == "NULL"
                                    else '"' + row[4][3] + '"',
                                    row[5],
                                    total,
                                )
                            print(query)
                            cursor.execute(
                                query
                            )
                            db.commit()
                            cursor.reset()
                        except Exception as e:
                            if "Duplicate entry" not in str(e):
                                print(query)
                                print(row)
                                print(e)
                                return {
                                    "error": "Error parsing data on row "
                                    + str(i + 12)
                                    + " of EL sheet. Please check the data and try again. "
                                    + str(e)
                                }
            return {"status": "success"}
        return {"status": "failed"}


@app.route("/staff/get_sheet_names", methods=["POST"])
def get_sheet_names():
    if session["role"] != 4:
        return redirect("/login")
    uploaded_file = request.files["file"]
    if uploaded_file:
        # Save the uploaded file to a desired location
        sheet_names = pd.ExcelFile(uploaded_file).sheet_names
        return {"status": "success", "sheet_names": sheet_names}
    return {"status": "failed"}


def normalize_date(obj):
    for i in range(len(obj)):
        obj[i] = list(obj[i])
        for j in range(len(obj[i])):
            if isinstance(obj[i][j], str):
                print(obj[i][j])
                if obj[i][j][0] == "[":
                    obj[i][j] = eval(obj[i][j])
            if isinstance(obj[i][j], bytes):
                obj[i][j] = ast.literal_eval(obj[i][j].decode("utf-8"))
            if isinstance(obj[i][j], datetime.date):
                obj[i][j] = obj[i][j].strftime("%d-%m-%Y")
    return obj


@app.route("/el/<id>", methods=["GET", "POST"])
def display_el(id):
    result = get_data(id, "el")
    if request.method == "POST":
        return json.dumps(normalize_date(result))
    return render_template(
        "table.html",
        table_content=enumerate(result),
        table_head=[
            "Si.No",
            "ID",
            "Department",
            "From",
            "To",
            "Prefix",
            "Suffix",
            "With permission on",
            "Date of join",
            "Total",
            "Document",
        ],
        type="Earned Leave (EL)",
        len=len(result),
    )


@app.route("/ml/<id>", methods=["GET", "POST"])
def display_ml(id):
    result = get_data(id, "ml")
    if request.method == "POST":
        return json.dumps(normalize_date(result))
    return render_template(
        "table.html",
        table_content=enumerate(result),
        table_head=[
            "Si.No",
            "ID",
            "Department",
            "From",
            "To",
            "Prefix",
            "Suffix",
            "Medical fitness on",
            "With permission on",
            "Date of join",
            "Total",
            "Document",
        ],
        type="Medical Leave (ML)",
        len=len(result),
    )


@app.route("/mtl/<id>", methods=["GET", "POST"])
def display_mtl(id):
    result = get_data(id, "mtl")
    if request.method == "POST":
        return json.dumps(normalize_date(result))
    return render_template(
        "table.html",
        table_content=enumerate(result),
        table_head=[
            "Si.No",
            "ID",
            "Department",
            "From",
            "To",
            "Prefix",
            "Suffix",
            "Medical fitness on",
            "With permission on",
            "Date of join",
            "Total",
            "Document",
        ],
        type="Maternal Leave (MTL)",
        len=len(result),
    )


@app.route("/lop/<id>", methods=["GET", "POST"])
def display_lop(id):
    result = get_data(id, "lop")
    if request.method == "POST":
        return json.dumps(normalize_date(result))
    return render_template(
        "table.html",
        table_content=enumerate(result),
        table_head=[
            "Si.No",
            "ID",
            "Department",
            "From",
            "To",
            "Prefix",
            "Suffix",
            "Medical fitness on",
            "With permission on",
            "Date of join",
            "Total",
            "Document",
        ],
        type="Loss Of Pay (LOP)",
        len=len(result),
    )


def getExtension(filename):
    return os.path.splitext(filename)[1]


@app.route("/upload_document/<dept>/<id>/<type>/<l_id>", methods=["GET", "POST"])
def upload_document(dept, id, type, l_id):
    if session["role"] != 4:
        return redirect("/login")
    if request.method == "POST":
        uploaded_files = request.files.getlist("files")
        labels = request.form.getlist("fileNames")
        upload_folder = os.path.join(
            app.config["DRIVE_DIRECTORIES"][dept], id, type, l_id
        )
        print(upload_folder)
        labels[0] = "leave_form"
        os.makedirs(upload_folder, exist_ok=True)
        for i in range(len(uploaded_files)):
            uploaded_file = uploaded_files[i]
            label = labels[i]
            if uploaded_file:
                uploaded_file.save(
                    os.path.join(
                        upload_folder, label + getExtension(uploaded_file.filename)
                    )
                )
            else:
                flash("Invalid file format", "error")
        return {"status": "success"}
    return "This is upload Page"


# Gets VL data, calculates vacation prevented dates and total no. of days availed
@app.route("/upload_vl/<id>/<sheet>", methods=["GET", "POST"])
def upload_vl_details(id, sheet):
    if session["role"] != 4:
        return redirect("/login")
    if request.method == "POST":
        uploaded_file = request.files["file"]
        if uploaded_file:
            vl_data = get_vl(id, uploaded_file, sheet, cursor)
            for key, value in vl_data.items():
                vl_data[key] = {
                    inner_key: convert_to_date_text(inner_value)
                    for inner_key, inner_value in value.items()
                }
            return {"data": vl_data, "status": "success"}
        else:
            return {"status": "failed"}
    return "This is VL data page"


@app.route("/update_vl/<staff_id>", methods=["GET", "POST"])
def update_vl(staff_id):
    if request.method == "POST":
        data = json.loads(request.data)
        for key, value in data.items():
            data[key] = {
                inner_key: convert_to_date_text(inner_value)
                for inner_key, inner_value in value.items()
            }
        print(staff_id)
        try:
            for i in data:
                data[i]["Availed_from"] = list(filter(lambda x: x not in [None, "NULL"], data[i]["Availed_from"]))
                data[i]["Availed_to"] = list(filter(lambda x: x not in [None, "NULL"], data[i]["Availed_to"]))
                query = (
                    f'INSERT into vl (vac_id, staff_id, availed_from, availed_to, prevented, other_leave) values ("%s", %s, %s, %s, %s, %s)'
                    % (
                        i,
                        staff_id,
                        "NULL"
                        if (
                            not data[i]["Availed_from"]
                        )
                        else '"' + str(data[i]["Availed_from"]) + '"',
                        "NULL"
                        if (
                            not data[i]["Availed_to"]
                        )
                        else '"' + str(data[i]["Availed_to"]) + '"',
                        "NULL"
                        if (
                            data[i]["Prevented"] is None
                            or data[i]["Prevented"] == "NULL"
                        )
                        else '"' + str(data[i]["Prevented"]) + '"',
                        "NULL"
                        if (
                                data[i]["others"] is None
                                or data[i]["others"] == "NULL"
                        )
                        else '"' + str(data[i]["others"]) + '"'
                    )
                )
                cursor.execute(query)

                db.commit()
        except Exception as e:
            print(e)
            return {"status": "failed", "error": str(e)}
        return {"status": "success"}
    return "This is VL data page"


# Sharon Work Starts
@app.route("/login", methods=["GET", "POST"])
def login():
    signined = None
    if request.method == "GET":
        return render_template("login.html", signined="")
    elif request.method == "POST":
        data = request.form
        mail_id = data["mail_id"]
        password = data["password"]

        if authenticate_user(mail_id, password):
            session_key = generate_session_key()
            update_session_key(mail_id, session_key)
            session["session_key"] = session_key
            signined = None
            return render_template("login.html", signined="true")
        else:
            return render_template("login.html", signined="false")


def authenticate_user(mail_id, password):
    hashed_password = hashlib.sha256(password.encode()).hexdigest()
    print(hashed_password)

    select_query = "SELECT id FROM user_login WHERE mail_id = %s AND password = %s"
    cursor.execute(select_query, (mail_id, hashed_password))
    user_id = cursor.fetchone()
    if user_id:
        return True
    else:
        return False


def generate_session_key():
    return "".join(random.choices(string.ascii_letters + string.digits, k=24))


def update_session_key(mail_id, session_key):
    update_query = "UPDATE user_login SET session_key = %s WHERE mail_id = %s"
    cursor.execute(update_query, (session_key, mail_id))
    db.commit()


@app.route("/forget-password", methods=["GET", "POST"])
def forgetPassword():
    signined = None
    if request.method == "GET":
        return render_template("forget_password.html")
    elif request.method == "POST":
        data = request.form
        mail_id = data["mail"]
        if check_email_in_database(mail_id):
            otp = str(otp_gen())
            session["otp"] = otp
            session["email_forgot"] = mail_id
            send_html_email(mail_id, otp, "Users")
            return render_template("one_time_password.html", otpf="")
        else:
            return render_template("forget_password.html", emailf="false")


def check_email_in_database(email):
    query = "SELECT COUNT(*) FROM user_login WHERE mail_id = %s"
    cursor.execute(query, (email,))
    result = cursor.fetchone()[0]

    if result > 0:
        return True
    else:
        return False


def otp_gen():
    return random.randint(100000, 999999)


def send_html_email(receiver_email, otp, name):
    sender_email = "sharonjoseph52@yahoo.com"
    password = "zgodamlpduxhfnrf"
    subject = "Faculty Login Password Change OTP"

    date = datetime.datetime.now().strftime("%B %d, %Y")

    html_content_1 = """
                    <!DOCTYPE html>
                    <html lang="en">
                    <head>
                        <meta charset="UTF-8" />
                        <meta name="viewport" content="width=device-width, initial-scale=1.0" />
                        <meta http-equiv="X-UA-Compatible" content="ie=edge" />
                        <title>Static Template</title>
                        <link
                        href="https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600&display=swap"
                        rel="stylesheet"
                        />
                    </head>
                    <body
                        style="margin: 0; font-family: 'Poppins', sans-serif; background: #ffffff; font-size: 14px;">
                        <div style="max-width: 680px; margin: 0 auto; padding: 45px 30px 60px; background: #f4f7ff; background-image: url(https://archisketch-resources.s3.ap-northeast-2.amazonaws.com/vrstyler/1661497957196_595865/email-template-background-banner); background-repeat: no-repeat; background-size: 800px 452px; background-position: top center; font-size: 14px; color: #434343;">
                        <header>
                            <table style="width: 100%;">
                            <tbody>
                                <tr style="height: 0;">
                                <td>
                                    <img src="https://upload.wikimedia.org/wikipedia/en/thumb/4/49/Anna_University_Logo.svg/330px-Anna_University_Logo.svg.png" height="100px"/>
                                    <h2 style="color: white;">
                                    Anna University Regional Campus - Tirunelveli
                                    </h2>
                                </td>
                                <td style="text-align: right;">
                                    <span style="font-size: 16px; line-height: 30px; color: #ffffff;">
                    """
    html_content_2 = """
                                    </span>
                                </td>
                                </tr>
                            </tbody>
                            </table>
                        </header>
                        <main>
                            <div style=" margin: 0; margin-top: 70px; padding: 92px 30px 115px; background: #ffffff; border-radius: 30px; text-align: center;">
                            <div style="width: 100%; max-width: 489px; margin: 0 auto;">
                                <h1 style=" margin: 0; font-size: 24px; font-weight: 500; color: #1f1f1f;">
                                Your OTP
                                </h1>
                                <p style=" margin: 0; margin-top: 17px; font-size: 16px; font-weight: 500;">
                                Hey """

    html_content_3 = """,
                </p>
                <p style=" margin: 0;  margin-top: 17px; font-weight: 500; letter-spacing: 0.56px;">
                Use the following OTP to complete the procedure to change your
                Password for Faculty Login. OTP is valid for
                <span style="font-weight: 600; color: #1f1f1f;">5 minutes</span>.
                Do not share this code with others.
                </p>
                <p style=" margin: 0; margin-top: 60px; font-size: 40px; font-weight: 600; letter-spacing: 15px; color: #ba3d4f;">
                """
    html_content_4 = """</p>
            </div>
            </div>
        </main>
        </div>
    </body>
    </html>"""

    html_content = (
        html_content_1
        + date
        + html_content_2
        + name
        + html_content_3
        + otp
        + html_content_4
    )

    # Create a multipart message
    message = MIMEMultipart()
    message["From"] = sender_email
    message["To"] = receiver_email
    message["Subject"] = subject

    # Attach HTML content to the email
    html_part = MIMEText(html_content, "html")
    message.attach(html_part)

    try:
        # Connect to the SMTP server
        with smtplib.SMTP("smtp.mail.yahoo.com", 587) as server:
            server.starttls()
            server.login(sender_email, password)
            server.sendmail(sender_email, receiver_email, message.as_string())
            return True

    except Exception as e:
        return False


@app.route("/otpr", methods=["POST"])
def otpr():
    return render_template("one_time_password.html", otpf="")


@app.route("/otp", methods=["POST"])
def otp():
    formotp = request.form["combinedValue"]
    sendotp = session.get("otp")
    if sendotp == formotp:
        session.pop("otp", None)
        return render_template("newpassword.html")
    else:
        return render_template("one_time_password.html", otpf="false")


@app.route("/newpass", methods=["POST"])
def newpass():
    passwd = request.form["password-2"]
    emaill = session.get("email_forgot")
    hashed_password = hashlib.sha256(passwd.encode()).hexdigest()
    update_password(emaill, hashed_password)
    session.pop("email_forgot", None)
    return render_template("newpassword.html", succ=True)


def update_password(email_to_search, new_password):
    cursor.execute(
        "UPDATE user_login SET password = %s WHERE mail_id = %s",
        (new_password, email_to_search),
    )


@app.route("/auth-success")
def auth_success():
    session_key = session.get("session_key")
    cursor.execute("SELECT role FROM user_login WHERE session_key = %s", (session_key,))
    session["role"] = cursor.fetchone()[0]
    print(session["role"])
    if session["role"] == 4 or session["role"] == 0:  # admin
        return redirect("/dashboard")
    else:
        return "You are not authorized to access this page"

# Auth Ends

def get_staff_from_name(name):
    cursor.execute(f'''
    SELECT id FROM staff
    WHERE name = '{name}'
    ''')
    namee = cursor.fetchone()
    cursor.reset()
    return jsonify(namee)

def set_run_font(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = font_size

def set_paragraph_alignment(paragraph, alignment):
    paragraph.alignment = alignment

# def replace_text_in_docx(input_docx_file, output_docx_file, replacements, data_array):
#     doc = docx.Document(input_docx_file)

#     # Replace text in paragraphs
#     for paragraph in doc.paragraphs:
#         for old_text, new_text in replacements.items():
#             if old_text in paragraph.text:
#                 new_paragraph = paragraph._element
#                 new_paragraph.clear_content()

#                 run = new_paragraph.add_run(new_text)
#                 set_run_font(run, "Arial", Pt(12))  # Set English font and size
#                 set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)
#             else:
#                 for run in paragraph.runs:
#                     set_run_font(run, "TamilFont", Pt(12))  # Set Tamil font and size

#     # Retrieve the first table in the document
#     table = doc.tables[0]

#     # Iterate through data array and add a new row to the table
#     for row_data in data_array:
#         new_row = table.add_row().cells
#         for i, cell_data in enumerate(row_data):
#             if i < len(new_row):
#                 cell = new_row[i]
#                 cell.text = str(cell_data)
#                 for paragraph in cell.paragraphs:
#                     set_run_font(paragraph.runs[0], "Bamini", Pt(12))  # Set Tamil font and size

#     doc.save(output_docx_file)

def replace_text_in_docx(input_docx_file, output_docx_file, replacements, data_array):
    doc = docx.Document(input_docx_file)
    if len(doc.tables) > 0:
        table = doc.tables[0]
        for row_data in data_array:
            new_row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(new_row):
                    cell = new_row[i]
                    cell.text = str(cell_data)
                    for paragraph in cell.paragraphs:
                        set_run_font(paragraph.runs[0], "Arial", Pt(10))
                        set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)
    for field, value in replacements.items():
        replace_placeholder(doc, field, value)
    doc.save(output_docx_file)





def replace_placeholder(doc, target, replacement):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if target in run.text and replacement:
                new_text = run.text.replace(target, replacement)
                run.text = new_text




def generate_random_string(length=5):
    characters = string.ascii_letters + string.digits 
    random_string = ''.join(random.choice(characters) for _ in range(length))
    return random_string

def insert_into_naan_mudhalvan_day(data):
    insert_query = """
    INSERT INTO naan_mudhalvan_day (department, semester, academic_year, academic_semester, day)
    VALUES (%s, %s, %s, %s, %s)
    """
    cursor.execute(insert_query, data)
    cursor.reset()
    db.commit()

def insert_into_timetable(data):
    insert_query = """
    INSERT INTO timetable (day, period, subject_id, academic_year, academic_semester)
    VALUES (%s, %s, %s, %s, %s)
    """
    
    cursor.executemany(insert_query, data)
    cursor.reset()
    db.commit()

def retrive_class_for_staff_in_specific_day(staff, day):
    query = f"""
                    SELECT tt.period, s.id, s.semester, s.department              
                    FROM timetable t
                    INNER JOIN timetable_timing tt ON t.period = tt.period
                    INNER JOIN subjects s ON t.subject_id = s.id
                    INNER JOIN staff_handling sh ON s.id = sh.subject_id
                    INNER JOIN staff st ON sh.staff_id = st.id
                    WHERE st.id = {staff} AND t.day = '{day}';
                """
    
    cursor.execute(query)
    classes = cursor.fetchall()
    return classes

def get_weekdays_in_date_range(start_date_str, end_date_str):
    weekdays = ["MON", "TUE", "WED", "THU", "FRI"]
    start_date = datetime.datetime.strptime(start_date_str, "%d-%m-%Y")
    end_date = datetime.datetime.strptime(end_date_str, "%d-%m-%Y")
    current_date = start_date
    result = []

    while current_date <= end_date:
        if current_date.weekday() < 5:
            day_of_week = weekdays[current_date.weekday()]
            result.append((day_of_week, current_date.strftime("%d-%m-%Y")))

        current_date += datetime.timedelta(days=1)

    return result


@app.route('/api/add_timetable', methods=['POST'])
def api_add_timetable():
    dep = request.form['dep']
    sem = request.form['sem']
    acayear = request.form['acayear']
    acasem = request.form['acasem']
    naan = request.form['naan']
    
    naan_mudhalvan_data = (dep, sem, acayear, acasem, naan)
    
    timetable_data = []
    for day in ['mon', 'tue', 'wed', 'thu', 'fri']:
        for period in range(1, 9):
            subject_id = request.form[f'{day}_{period}']
            if subject_id == '000000':
                continue
            else:
                timetable_data.append((day.upper(), period, subject_id, acayear, acasem))
    
    insert_into_naan_mudhalvan_day(naan_mudhalvan_data)
    insert_into_timetable(timetable_data)
    
    return "Hello"

@app.route('/add_timetable', methods=['GET'])
def add_timetable():
    return render_template('add_timetable.html')


@app.route('/api/update_timetable_timing', methods=['POST'])
def update_timetable_timing():         
    for period in range(1, 9):
        start_time = request.form.get(f'start_{period}')
        end_time = request.form.get(f'end_{period}')
                
        update_query = f"UPDATE `timetable_timing` SET `start_time`='{start_time}',`end_time`='{end_time} WHERE  `period`='{period}''"
        #update_query = f"INSERT INTO `timetable_timing`(`start_time`, `end_time`, `period`) VALUES ('{start_time}','{end_time}','{period}')"
        cursor.execute(update_query)
        cursor.reset()
        db.commit()    
    return "True"

@app.route('/update_timetable_timing', methods=['GET'])
def update_timetable_timing_():          
    return render_template('update_timetable_timing.html')

@app.route('/api/insert_subjects', methods=['POST'])
def insert_subject():
    code = request.form.get("code")
    name = request.form.get("name")
    type = request.form.get("type")
    sem = request.form.get("sem")
    credit = request.form.get("credit")
    contact = request.form.get("contact")
    regu = request.form.get("regu")
    department = request.form.get("department")
    staff = request.form.get("staff")

    update_query = f"""
     INSERT 
     INTO `subjects`(`id`, `name`, `type`, `semester`, `credit_hour`, `contact_period`, `regulation`, `department`) 
     VALUES ('{code}', '{name}', '{type}', '{sem}', '{credit}', '{contact}', '{regu}', '{department}');
    """

    cursor.execute(update_query)
    cursor.reset()

    update_query = f"INSERT INTO `staff_handling`(`subject_id`, `staff_id`) VALUES ('{code}','{staff}')"
    cursor.execute(update_query)
    cursor.reset()
    db.commit()

    return "Subject updated successfully!"

@app.route('/add_subjects', methods=['GET'])
def add_subject():
    return render_template('add_subject.html')

@app.route("/api/retrive_all_staff", methods=["GET"])
def retrive_all_staff():
    cursor.execute('SELECT `id`,`name` FROM `staff`')
    staffs = cursor.fetchall()
    return jsonify(staffs)

@app.route("/api/retrive_all_subjects", methods=["GET"])
def retrive_all_subjects():
    cursor.execute('SELECT `id`,`name` FROM `subjects`')
    subjects = cursor.fetchall()
    return jsonify(subjects)


@app.route("/api/get_staffs_classes/<staff>/<fdate>/<tdate>", methods=["GET"])
def get_staffs_classes(staff, fdate, tdate):
    days = get_weekdays_in_date_range(fdate, tdate)
    staffs_classes = []
    
    for day in days:
        day_name, date = day[0], day[1]
        classes = retrive_class_for_staff_in_specific_day(staff, day_name)
        for clas in classes:
            staffs_classes.append([day_name, date, clas[0], clas[1], clas[2], clas[3]])
    return jsonify(staffs_classes)


@app.route("/api/get_free_staffs/<day>/<period>/<semester>/<department>", methods=["GET"])
def get_free_staffs(day, period, semester, department):
    cursor.execute('''SELECT
    tt.day AS timetable_day,
    tt.period AS timetable_period,
    s.id AS subject_id,
    s.name AS subject_name,
    s.semester AS subject_semester,
    s.department AS subject_department,
    sh.staff_id,
    st.name AS staff_name
FROM
    timetable tt
JOIN
    subjects s ON tt.subject_id = s.id
JOIN
    staff_handling sh ON tt.subject_id = sh.subject_id
JOIN
    staff st ON sh.staff_id = st.id;''')

    data = get_free_staff(cursor.fetchall(), str(department), str(day), str(period), str(semester))

    cursor.reset()
    return jsonify(data)

def get_free_staff(timetable, department, day, period, semester):
    staff_with_class = set()
    for row in timetable:
        if str(row[0]) == day and str(row[1]) == period:
            staff_with_class.add((row[6], row[7]))
      
    staff_in_department = set()
    for row in timetable:
        if str(row[5]) == department and str(row[4]) == semester:
            staff_in_department.add((row[6], row[7]))

    free_staff = staff_in_department - staff_with_class

    return list(free_staff)


@app.route('/api/get_staffs_subjects/<staff>/<department>/<semester>', methods=['GET'])
def get_staffs_subjects(staff, department, semester):
    query = f"""SELECT
    s.id AS subject_code,
    s.name AS subject_name
FROM
    staff_handling sh
JOIN
    subjects s ON sh.subject_id = s.id
JOIN
    staff st ON sh.staff_id = st.id
WHERE
    st.id = {staff}
    AND s.department = '{department}'
    AND s.semester = {semester}"""

    cursor.execute(query)

    data = cursor.fetchall()

    cursor.reset()

    return jsonify(data)
    

@app.route("/api/get_staff_id_from_name/<name>", methods=["GET"])
def get_staff_id_from_name(name):
    cursor.execute(f'''
    SELECT id FROM staff
    WHERE name = '{name}'
    ''')

    namee = cursor.fetchone()

    cursor.reset()
    return jsonify(namee)


@app.route("/api/changeStaff/<department>/<name>", methods=["GET"])
def changeStaff(department, name):
    query = f'''SELECT staff_handling.subject_id, subjects.name
FROM staff_handling
JOIN subjects ON subjects.id = staff_handling.subject_id
JOIN staff ON staff.id = staff_handling.staff_id
WHERE staff.name = '{name}'
AND subjects.department = '{department}';'''
    cursor.execute(query)
    subjects = cursor.fetchall()
    cursor.reset()
    return jsonify(subjects)

@app.route("/api/generate/el", methods=["POST"])
def genRate_el():
    data = request.json
    object_data = data.get('objectData', {})
    array_data = data.get('leaves', []) 
    name = generate_random_string()
    name = name + '.docx'
    temp_dir = 'temp/'
    os.makedirs(temp_dir, exist_ok=True)
    file_path = os.path.join(temp_dir, name)
    replace_text_in_docx('leaveFormTemplates/el.docx', file_path, object_data, array_data)
    convert(file_path, file_path.replace('.docx', '.pdf'))
    return jsonify(file_path.replace('.docx', '.pdf'))

@app.route("/api/generate/cl", methods=["POST"])
def genRate_cl():
    data = request.json
    object_data = data.get('objectData', {})
    array_data = data.get('leaves', []) 
    name = generate_random_string()
    name = name + '.docx'
    temp_dir = 'temp/'
    os.makedirs(temp_dir, exist_ok=True)
    file_path = os.path.join(temp_dir, name)
    replace_text_in_docx('leaveFormTemplates/cl.docx', file_path, object_data, array_data)
    convert(file_path, file_path.replace('.docx', '.pdf'))
    os.remove(file_path)
    return jsonify(file_path.replace('.docx', '.pdf'))

@app.route('/leave_form', methods=['GET'])
def leave_form():
    return render_template('leave_form.html')

@app.route('/api/download/<folder>/<file>', methods=['GET'])
def download(folder, file):
    file = f"{folder}/{file}"
    return send_file(file, as_attachment=True)

# Ended Sharon Work


# Maheshwar work starts here


@app.route(
    "/generate_report/<staff_id>/<need_attendance>/<remarks>/<need_pdf>",
    methods=["POST", "GET"],
)
def report_generation(staff_id, need_attendance, remarks, need_pdf):
    if request.method == "GET":
        remarks = urllib.parse.unquote(remarks)
        remarks = remarks if remarks != "null" else ""
        file = generate_report(
            cursor,
            staff_id,
            True if need_attendance == "1" else False,
            remarks,
            True if need_pdf == "1" else False,
        )
        try:
            return send_file(
                "static/reports/" + file["name"] + "." + file["type"],
                as_attachment=True,
            )
        except Exception as e:
            print("Error:", str(e))
            return {"status": "failed", "error": str(e)}

    return render_template("report_generation.html")


@app.route("/get_leave_with_date/", methods=["POST"])
def get_dates_with_range():
    if request.method == "POST":
        start = request.form.get("start")
        end = request.form.get("end")
        staff_id = request.form.get("id")
        query = (
            f"SELECT * FROM `el` WHERE `from` BETWEEN '%s 00:00:00' AND '%s 00:00:00' AND id=%s and `to` BETWEEN '%s 00:00:00' AND '%s 00:00:00';"
            % (start, end, staff_id, start, end)
        )
        print(query)
        cursor.execute(query)
        data = cursor.fetchall()
        return json.dumps(normalize_date(data))


@app.route("/get_files/", methods=["GET", "POST"])
def get_files():
    if request.method == "POST":
        path_parts = request.form.get("path").split("_")
        print(path_parts)
        department = path_parts[0]
        staff_id = path_parts[1]
        leave_type = path_parts[2]
        leave_id = path_parts[3]
        # path = os.path.join(drive_folder,department, staff_id, leave_type, leave_id)
        path = os.path.join(
            "H:", "My Drive", "data", department, staff_id, leave_type, leave_id
        )
        print(path)
        if os.path.exists(path):
            print("Path exists")
            files = os.listdir(path)
            print(files)
            return jsonify({"files": files})
        else:
            print("Path does not exist")
            return jsonify({"status": "failed", "message": "No files exist"})
        # print(os.path.exists(path))
        # if os.path.exists(path):
        #     files = os.listdir(path)
        #     print(files)
        #     return jsonify({"files": files})
        # else:
        #     return jsonify({"message": "No files exist"})

@app.route("/get_files/<path:path>", methods=["GET", "POST"])
def get_file(path):
    if request.method == "GET":
        file_path = os.path.join("H:", "My Drive", "data", path)
        if os.path.exists(file_path):
            print("Path exists", file_path)
            try:
                return send_file(file_path, as_attachment=False)
            except PermissionError:
                return jsonify({"status": "failed", "message": "Permission denied"})
        else:
            return jsonify({"status": "failed", "message": "File does not exist"})

@app.route('/get_el_earned/', methods=["POST"])
def send_el_earned():
    if request.method == "POST":
        id = request.form.get('id')
        bef_probation, after_probation = get_el_earned(id, cursor)
        bef_probation['working_from'] = bef_probation['working_from'].apply(lambda x: x.strftime("%Y-%m-%d"))
        bef_probation['working_to'] = bef_probation['working_to'].apply(lambda x: x.strftime("%Y-%m-%d"))
        after_probation['working_from'] = after_probation['working_from'].apply(lambda x: x.strftime("%Y-%m-%d"))
        after_probation['working_to'] = after_probation['working_to'].apply(lambda x: x.strftime("%Y-%m-%d"))
        # print(bef_probation.T.to_dict())
        return jsonify({"before": bef_probation.T.to_dict(), "after": after_probation.T.to_dict()})

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0")
